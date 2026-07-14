#!/usr/bin/env python3
"""Rebuild reference.docx and reference-report.docx from TenX brand
colors/fonts/logo.

Run this whenever colors.md, typography.md, or the logo change. Produces
two Pandoc --reference-doc templates, sharing the same fonts/headings/
footer machinery but with two different title-page treatments:

  - reference.docx — policies and templates: a restrained, typographic
    title page (see build_policy_reference()).
  - reference-report.docx — reports (quarterly business reports and
    similar): a graphic cover page (see build_report_reference()) — the
    "annual report / pitch deck" genre the policy title page deliberately
    avoids (see the design-history comment below).

Requires: pandoc, python-docx (`pip install python-docx`).
"""
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).resolve().parent.parent
PANDOC_DIR = Path(__file__).resolve().parent
PREVIEW_CONTENT_POLICY = PANDOC_DIR / "preview-content.md"
PREVIEW_CONTENT_REPORT = PANDOC_DIR / "preview-content-report.md"
POLICY_OUTPUT = PANDOC_DIR / "reference.docx"
REPORT_OUTPUT = PANDOC_DIR / "reference-report.docx"

# Stacked lockup (icon above wordmark), not the horizontal one — per direct
# feedback that the mark should read icon-over-text, not icon-beside-text.
# 512px is far more resolution than either print size below needs, but
# logo/build-logos.py's smallest common size (128px) would be visibly soft,
# so round up rather than generate one-off in-between sizes.
LOGO_BLACK = REPO_ROOT / "logo" / "png" / "lockup-stacked-black-512.png"
LOGO_WHITE = REPO_ROOT / "logo" / "png" / "lockup-stacked-white-512.png"

ACCENT = RGBColor(0xA8, 0x55, 0xF7)   # colors.md: Accent (violet)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)     # colors.md: Gray — mid
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COVER_HEX = "3D1A5B"                  # colors.md: MASTHEAD_HEX — deep violet, reused here as the report cover's background
FONT = "Calibri"                      # typography.md

LOGO_HEIGHT = Cm(1.5)        # restrained header size — every policy/template page, and a report's page 2+
COVER_LOGO_HEIGHT = Cm(3.2)  # report cover only (page 1) — large and centered, the graphic focal point

# Title-page design history (reference.docx — policies/templates), oldest first:
# 1. A curve+badge cover graphic (a large violet sweep + circular logo
#    badge) was built and rendered, then dropped after review — real SOC 2
#    / ISO 27001 policy documents converge on restrained, typographic title
#    pages, and illustrative shapes at fixed positions can't adapt to a
#    variable-length title without either colliding with the text or
#    requiring per-document layout.
# 2. A shaded "masthead band" (Title as a colored panel, inset within the
#    margins) replaced it — paragraph shading grows with wrapped text
#    automatically, a fixed shape does not.
# 3. The masthead band itself was later dropped per direct feedback: no
#    background color on Title, and the TenX logo moved from a small
#    top-left header mark to a large, centered lockup directly above the
#    title, so logo + title + Author/Date now read as one cohesive,
#    centered unit instead of a colored box with a corner logo.
#
# That graphic-cover genre wasn't wrong for every document type, just wrong
# for restrained policy documents — reports (quarterly business reports and
# similar) are exactly the genre it fits. build_report_reference() below
# revives the same shaded-panel technique from step 2, scaled up, using
# COVER_HEX and the white logo variant so the mark reads against the panel
# instead of against the page.


def set_style_font_color(doc, style_name, color=None, bold=None, size=None):
    style = doc.styles[style_name]
    style.font.name = FONT
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), FONT)
    # Pandoc's default styles carry theme-font attributes (w:asciiTheme etc.)
    # on the same rFonts element, and per ECMA-376 §17.3.2.26 theme
    # attributes OVERRIDE the literal ones when both are present — leaving
    # them in place means the style silently renders in the document theme's
    # font (Aptos Display in Pandoc 3.x), not Calibri. Strip them so the
    # literal values actually win.
    for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if size is not None:
        style.font.size = size


def style_title(doc, size, text_color, bold=True, space_before=Pt(18), space_after=Pt(16),
                 shading_hex=None, pad_h_twips=0):
    """Title style. With shading_hex=None (policies/templates): plain
    colored text, no box — the logo above it in the header carries the
    visual weight instead (see build_policy_reference()). With
    shading_hex set (reports): a colored panel, same mechanic as the
    dropped policy masthead band — paragraph shading covers the
    paragraph's own space_before/space_after too, so those act as the
    panel's internal top/bottom padding, not just inter-paragraph gaps.

    Pandoc's default Title is already center-aligned, so centering
    doesn't need to be set here."""
    style = doc.styles['Title']
    pPr = style.element.get_or_add_pPr()
    if shading_hex is not None:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), shading_hex)
        pPr.append(shd)
    if pad_h_twips:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(pad_h_twips))
        ind.set(qn('w:right'), str(pad_h_twips))
        pPr.append(ind)
    style.font.color.rgb = text_color
    style.font.size = size
    style.font.bold = bold
    style.paragraph_format.space_before = space_before
    style.paragraph_format.space_after = space_after
    style.paragraph_format.line_spacing = 1.15


def style_meta_line(doc, style_name, size, text_color, space_before=Pt(0), space_after=Pt(2)):
    """Author/Date: plain text below the title, color/size/weight only.
    Shading/indent are deliberately left unset here — Author/Date are
    basedOn Title in Pandoc's default reference.docx, so whatever Title
    carries (nothing, for policies; COVER_HEX shading, for reports) is
    inherited automatically through that cascade. For the report cover,
    that's exactly what's wanted: Author/Date continue the same colored
    panel Title started, using their own space_before/space_after as
    padding the same way Title's does."""
    style = doc.styles[style_name]
    style.font.color.rgb = text_color
    style.font.size = size
    style.font.bold = False
    style.paragraph_format.space_before = space_before
    style.paragraph_format.space_after = space_after


def set_header(header_obj, logo_png, logo_height, alignment=1, shading_hex=None,
                space_before=Pt(0), space_after=Pt(0)):
    header_obj.is_linked_to_previous = False
    para = header_obj.paragraphs[0]
    if shading_hex is not None:
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), shading_hex)
        pPr.append(shd)
    para.alignment = alignment
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after = space_after
    run = para.add_run()
    run.add_picture(str(logo_png), height=logo_height)


def add_footer_page_field(section):
    """Centered page number — but not on page 1 (always the title/cover
    page, or the only page for a short template).

    The XML this produces is spec-correct (titlePg set, a distinct empty
    w:type="first" footer, a separate w:type="default" footer carrying
    the PAGE field) and Word/Google Docs honor it exactly as intended.
    LibreOffice 26.2 headless does not: verified with an isolated minimal
    docx that as soon as ANY header is present — even one identically
    defined for first-page and default, with no other asymmetry at all —
    LO leaks the default footer's page number onto page 1 regardless.
    Same category of LO-headless-only limitation as the TOC-field-
    population issue noted below (a rendering tool gap, not a document
    defect) — Word and Google Docs are the real targets."""
    section.different_first_page_header_footer = True
    section.first_page_footer.is_linked_to_previous = False
    section.first_page_footer.paragraphs[0].text = ""

    section.footer.is_linked_to_previous = False
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = 1  # center
    footer_para.text = ""
    run = footer_para.add_run()
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = "PAGE"
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_end)


def base_document():
    """Build the shared base every reference doc starts from: Pandoc's
    own default reference.docx, with page geometry, fonts/heading colors,
    the page-number footer, and the update-fields flag applied. Each
    variant then layers its own title-page/header treatment and margins
    on top of this."""
    base_path = Path(tempfile.mkdtemp()) / "default-reference.docx"
    subprocess.run(
        ["pandoc", "-o", str(base_path), "--print-default-data-file", "reference.docx"],
        check=True,
    )
    doc = Document(str(base_path))

    # Pandoc's default reference.docx doesn't set page size/margins
    # explicitly (python-docx reads them back as None) — pin them so the
    # fixed vertical-positioning math each variant does is against known
    # numbers, not whatever the rendering application happens to default
    # to. top_margin/header_distance are set per-variant, not here — each
    # title-page treatment needs a different amount of header room.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    set_style_font_color(doc, 'Normal')
    set_style_font_color(doc, 'Heading 1', color=ACCENT, bold=True, size=Pt(18))
    set_style_font_color(doc, 'Heading 2', color=ACCENT, bold=True, size=Pt(14))
    set_style_font_color(doc, 'Heading 3', color=ACCENT, bold=True, size=Pt(12))
    set_style_font_color(doc, 'TOC Heading', color=ACCENT, bold=True, size=Pt(18))
    # Title/Author/Date need the same Calibri treatment: style_title and
    # style_meta_line set color/size/shading but never touch the font
    # name, so without these calls the three title-page styles keep ONLY
    # Pandoc's theme-font attributes and render in the theme font instead
    # of Calibri.
    set_style_font_color(doc, 'Title')
    set_style_font_color(doc, 'Author')
    set_style_font_color(doc, 'Date')

    # Page break before TOC Heading — separates the title/cover page from
    # the TOC page. Policies and reports pass --toc; templates don't, so
    # this only ever fires for the other two.
    doc.styles['TOC Heading'].paragraph_format.page_break_before = True

    add_footer_page_field(section)

    # Tell any compliant consumer (Word, in particular) to refresh all
    # fields — including the TOC field — the moment the document is
    # opened, so a reader never has to manually press F9. This does NOT
    # fix LibreOffice's headless PDF export, which was confirmed (by
    # testing both this flag and an export-filter option) to leave the
    # TOC field unpopulated regardless — a rendering-tool limitation, not
    # a document defect. Word honors this flag per spec, and Google Docs
    # independently converts the field to its own live, auto-refreshing
    # TOC object at import time either way.
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    doc.settings.element.insert(0, update_fields)

    return doc, section


def render(doc, preview_content, output):
    with tempfile.TemporaryDirectory() as tmp:
        # Pandoc's own default reference.docx body is just a swatch of
        # literal style-name placeholders ("Heading 1", "Body Text.", ...)
        # — useless for judging what a real document actually looks like.
        # Rendering a realistic preview document through this styled base
        # (via --reference-doc) bakes real, meaningful demo content into
        # the final reference doc while keeping every style, header, and
        # footer intact — pandoc copies those through unchanged from
        # whatever --reference-doc it's given.
        styled_base_path = Path(tmp) / "styled-base.docx"
        doc.save(str(styled_base_path))
        subprocess.run(
            [
                "pandoc", str(preview_content),
                f"--reference-doc={styled_base_path}",
                "--toc",
                "-o", str(output),
            ],
            check=True,
        )
        print(f"saved {output}")


def build_policy_reference():
    """Policies and templates: restrained, typographic title page. Logo,
    centered and large enough to anchor the page, sits directly above
    plain colored Title text — no panel — so logo+title+Author/Date read
    as one cohesive block. Identical header on every page (not just the
    title page): with titlePg set (see add_footer_page_field), an
    undefined first-page header renders BLANK per spec rather than
    inheriting the default, and the title page is exactly where the logo
    needs to be — so it needs its own explicit reference regardless; a
    consistent brand mark on every page is also just a reasonable look on
    its own."""
    doc, section = base_document()

    # top_margin is section-wide — it also sets where body text starts on
    # every page after the title page (TOC, then real content), not just
    # page 1. Sized to fit header_distance + LOGO_HEIGHT plus a small
    # buffer, and kept as tight as that requires rather than pushed
    # further down: doing the latter to visually center the logo+title
    # block on the title page would waste that same extra space on every
    # subsequent page. Genuine vertical centering of a header-anchored
    # element is out of reach for that reason — same kind of hard OOXML
    # wall as rounded panel corners — so this settles for "logo sits
    # directly above the title as one tight, cohesive unit in the upper
    # part of the page" rather than "block is dead-center."
    section.header_distance = Inches(0.6)
    section.top_margin = Inches(1.3)

    style_title(doc, size=Pt(24), text_color=ACCENT, space_before=Pt(18), space_after=Pt(16))
    style_meta_line(doc, 'Author', size=Pt(11), text_color=GRAY, space_before=Pt(100), space_after=Pt(2))
    style_meta_line(doc, 'Date', size=Pt(10), text_color=GRAY, space_after=Pt(0))

    for header_obj in (section.first_page_header, section.header):
        set_header(header_obj, LOGO_BLACK, LOGO_HEIGHT, alignment=1)

    render(doc, PREVIEW_CONTENT_POLICY, POLICY_OUTPUT)


def build_report_reference():
    """Reports (quarterly business reports and similar): a graphic cover
    page. The header (page 1 only) and Title/Author/Date are all shaded
    the same COVER_HEX, with no unshaded gap between the header's shaded
    box and Title's — see style_title's docstring on why paragraph
    shading covering space_before/space_after makes that a matter of
    sizing, not a separate mechanism. Interior pages (2+) fall back to
    the same small black header used by policies/templates — the graphic
    treatment is deliberately cover-only, not carried through the whole
    report."""
    doc, section = base_document()

    # header_distance + the cover header's own space_before/logo/space_after
    # need to land on top_margin with (ideally) zero gap, or the shaded
    # header box and shaded Title box show a visible unshaded seam between
    # them. Tuned empirically against the rendered preview, same as every
    # other spacing constant in this file — not a closed-form calculation.
    section.header_distance = Inches(0.4)
    section.top_margin = Inches(2.85)

    style_title(
        doc, size=Pt(34), text_color=WHITE, space_before=Pt(0), space_after=Pt(20),
        shading_hex=COVER_HEX, pad_h_twips=360,
    )
    style_meta_line(doc, 'Author', size=Pt(13), text_color=WHITE, space_before=Pt(0), space_after=Pt(4))
    # Date is the last paragraph in the shaded panel — its space_after is
    # the panel's bottom padding, same role Title's space_after/space_before
    # play above.
    style_meta_line(doc, 'Date', size=Pt(11), text_color=RGBColor(0xD8, 0xC9, 0xEE), space_after=Pt(48))

    set_header(
        section.first_page_header, LOGO_WHITE, COVER_LOGO_HEIGHT, alignment=1,
        shading_hex=COVER_HEX, space_before=Pt(40), space_after=Pt(24),
    )
    set_header(section.header, LOGO_BLACK, LOGO_HEIGHT, alignment=1)

    render(doc, PREVIEW_CONTENT_REPORT, REPORT_OUTPUT)


def main():
    build_policy_reference()
    build_report_reference()


if __name__ == "__main__":
    main()
